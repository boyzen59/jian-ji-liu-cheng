#!/usr/bin/env python3
"""Build the six planning/production deliverables from a project scaffold."""

from __future__ import annotations

import argparse
import json
import shutil
import sys
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError as exc:  # pragma: no cover
    raise SystemExit("python-docx is required. Use the bundled Codex document Python runtime.") from exc

from build_text_review_docx import build_document as build_text_review_document
from build_text_review_docx import unresolved_rows


ACCENT = RGBColor(232, 115, 74)
INK = RGBColor(61, 48, 40)
MUTED = RGBColor(112, 96, 85)
PAPER = "FFF8F0"
EXPECTED_OUTPUT = {
    "width": 2560,
    "height": 1440,
    "fps": 60,
    "frame_rate_mode": "CFR",
    "container": "mp4",
    "video_codec": "h264",
    "video_profile": "high",
    "pixel_format": "yuv420p",
    "color_primaries": "bt709",
    "color_transfer": "bt709",
    "color_matrix": "bt709",
    "audio_mode": "no_audio_stream",
    "burned_in_captions": False,
}


def load_json(path: Path) -> dict[str, Any]:
    try:
        return json.loads(path.read_text(encoding="utf-8-sig"))
    except FileNotFoundError as exc:
        raise SystemExit(f"Missing required project file: {path}") from exc
    except json.JSONDecodeError as exc:
        raise SystemExit(f"Invalid JSON {path}: {exc}") from exc


def value_text(value: Any) -> str:
    if value is None or value == "":
        return "—"
    if isinstance(value, bool):
        return "是" if value else "否"
    if isinstance(value, (list, tuple)):
        return "、".join(value_text(item) for item in value) if value else "—"
    if isinstance(value, dict):
        return json.dumps(value, ensure_ascii=False, sort_keys=True)
    return str(value)


def parse_timecode(value: Any) -> float:
    if isinstance(value, (int, float)):
        return float(value)
    parts = str(value).replace(",", ".").split(":")
    try:
        if len(parts) == 3:
            return int(parts[0]) * 3600 + int(parts[1]) * 60 + float(parts[2])
        if len(parts) == 2:
            return int(parts[0]) * 60 + float(parts[1])
        return float(parts[0])
    except (ValueError, IndexError) as exc:
        raise ValueError(f"invalid timecode: {value}") from exc


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)


def style_document(doc: Document, title: str, project: str, generated: datetime) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.15
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style.font.color.rgb = INK
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)
    header = section.header.paragraphs[0]
    header.text = f"{project}  |  H.264 · Rec.709 · 2560×1440 · 60fps · 静音母版"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = MUTED
    add_page_number(section.footer.paragraphs[0])
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = ACCENT
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(f"{project}  |  生成时间 {generated.isoformat(timespec='seconds')}")
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = MUTED


def add_label(doc: Document, label: str, value: Any) -> None:
    paragraph = doc.add_paragraph()
    label_run = paragraph.add_run(label + "  ")
    label_run.bold = True
    label_run.font.color.rgb = ACCENT
    value_run = paragraph.add_run(value_text(value))
    value_run.font.color.rgb = INK


def add_entry(doc: Document, heading: str, fields: Iterable[tuple[str, Any]]) -> None:
    doc.add_heading(heading, level=2)
    for label, value in fields:
        add_label(doc, label, value)


def save_document(doc: Document, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def validate_inputs(project: dict[str, Any], spec: dict[str, Any], assets: dict[str, Any]) -> list[str]:
    errors: list[str] = []
    output = spec.get("output", project.get("spec", {}))
    for key, value in EXPECTED_OUTPUT.items():
        if output.get(key) != value:
            errors.append(f"output.{key} must be {value!r}")
    scenes = spec.get("scenes", [])
    if not scenes:
        errors.append("video-spec scenes must not be empty")
    ai_videos = assets.get("ai_video_requests", [])
    if len(ai_videos) > 250:
        errors.append(f"ai_video_requests must not exceed 250; got {len(ai_videos)}")
    for item in ai_videos:
        try:
            duration = float(item.get("duration_seconds", 0))
        except (TypeError, ValueError):
            duration = 0
        if not 5.0 <= duration <= 15.0:
            errors.append(f"AI video {item.get('id', '?')} duration_seconds must be between 5.0 and 15.0")
        if not item.get("prompt_zh") or not item.get("prompt_en"):
            errors.append(f"AI video {item.get('id', '?')} requires prompt_zh and prompt_en")
        if item.get("no_subtitles") is not True:
            errors.append(f"AI video {item.get('id', '?')} no_subtitles must be true")
    for list_name in ["ai_image_requests"]:
        for item in assets.get(list_name, []):
            if not item.get("prompt_zh") or not item.get("prompt_en"):
                errors.append(f"AI image {item.get('id', '?')} requires prompt_zh and prompt_en")
            required = ["semantic_segment_id", "source_line", "timecode", "planned_duration_seconds", "variant_role"]
            missing = [key for key in required if item.get(key) in (None, "")]
            if missing:
                errors.append(f"AI image {item.get('id', '?')} missing coverage fields: {', '.join(missing)}")
            if item.get("expected_use_count", 1) != 1:
                errors.append(f"AI image {item.get('id', '?')} expected_use_count must be 1")
    return errors


def build_execution_manual(
    root: Path,
    project: dict[str, Any],
    brief: str,
    outline: dict[str, Any],
    spec: dict[str, Any],
    animation: dict[str, Any],
    assets: dict[str, Any],
    presenter: dict[str, Any],
    review: dict[str, Any],
    generated: datetime,
) -> str:
    lines = [
        "视频剪辑执行说明书",
        "=" * 24,
        f"项目：{project.get('project', '')}",
        f"版本：{spec.get('version', 'V1')}",
        f"生成时间：{generated.isoformat(timespec='seconds')}",
        "用途：把本文件交给一个新任务后，可在不读取原聊天的情况下继续生产。",
        "",
        "一、开始执行前",
        "1. 逐一检查下列输入文件、版本和来源；文件缺失或哈希变化时停止对应环节。",
        "2. 以锁定口播精确时长为总时间真本，以中文 SRT 为外挂字幕时间骨架；连续字幕不得烧录。",
        "3. 既有文字继承已通过基线；只有新增或改字进入《动效文字执行手册》增量门禁，纯版式调整不重审。",
        "4. 先完成场景类型大纲，为每章建立全屏超大双语标题总览，再按 video spec 实现；Remotion 是唯一总时间线。",
        "",
        "二、固定主交付",
        "MP4；H.264/AVC High；yuv420p；Rec.709（primaries/transfer/matrix 均 bt709）；2560×1440；60fps CFR；无音频流；无烧录连续字幕。禁止 H.265/HEVC。",
        "Remotion 主合成显式设 width=2560、height=1440、fps=60、codec=h264、muted，且不用 enforce-audio-track。最终 MP4 用 scripts/verify_master_media.py + ffprobe 验收。",
        "若需 FFmpeg 规范化（以下命令假定输入已是正确 Rec.709）：",
        'ffmpeg -i master_input.mp4 -map 0:v:0 -an -vf "scale=2560:1440:flags=lanczos:force_original_aspect_ratio=decrease,pad=2560:1440:(ow-iw)/2:(oh-ih)/2:color=black,format=yuv420p,setparams=range=tv:color_primaries=bt709:color_trc=bt709:colorspace=bt709" -r 60 -fps_mode cfr -c:v libx264 -profile:v high -level:v 5.1 -preset slow -crf 17 -pix_fmt yuv420p -color_range tv -color_primaries bt709 -color_trc bt709 -colorspace bt709 -movflags +faststart master_h264_rec709_2k60_silent.mp4',
        "警告：setparams/metadata 只写标签，不会把错误源色彩真正转换为 Rec.709。HDR、P3、BT.2020 或标签缺失素材必须先按真实输入做色彩转换并抽帧复核。",
        "",
        "三、固定视觉与真人规则",
        "Soft Signal / 亲密·温暖 / Apple-inspired frosted glass。文字完全不透明；正文对比度至少 4.5:1，大字至少 3:1；玻璃只包局部文字，图片/视频/证据图本体保持清晰且 blur=0。",
        "真人累计时长必须小于总时长20%；只有总开头和实际最终结尾全屏；正文固定使用抬高右下圆窗；最终全屏持续到最后一帧。说明动效逐镜比较Remotion/HyperFrames。",
        "",
        "四、项目简报",
        brief.strip() or "[待补充]",
        "",
        "五、逐镜执行",
    ]
    for scene in spec.get("scenes", []):
        lines.extend([
            f"[{scene.get('id', '?')}] {scene.get('start', '')}–{scene.get('end', '')}",
            f"章节：{value_text(scene.get('chapter'))}",
            f"原文范围：{value_text(scene.get('source_range'))}",
            f"场景类型：{value_text(scene.get('scene_type'))}",
            f"功能：{value_text(scene.get('function'))}",
            f"观众先看：{value_text(scene.get('viewer_should_notice'))}",
            f"看完理解：{value_text(scene.get('viewer_should_understand'))}",
            f"载体/引擎：{value_text(scene.get('carrier'))} / {value_text(scene.get('engine'))}",
            f"双引擎同级选型：{value_text(scene.get('engine_selection'))}",
            f"程序化机会：{value_text(scene.get('programmatic_opportunity'))} / {value_text(scene.get('programmatic_declined_reason'))}",
            f"构图：{value_text(scene.get('composition'))}",
            f"视觉页/内容指纹：{value_text(scene.get('visual_page_id'))} / {value_text(scene.get('content_fingerprint'))}",
            f"补时策略/内部阶段：{value_text(scene.get('duration_fill_strategy'))} / {value_text(scene.get('internal_phases'))}",
            f"屏显英文：{value_text(scene.get('screen_en'))}",
            f"屏显中文：{value_text(scene.get('screen_zh'))}",
            f"素材：{value_text(scene.get('assets'))}",
            f"语义段/文字来源：{value_text(scene.get('semantic_segment_id'))} / {value_text(scene.get('source_text_ids'))}",
            f"素材锁定指派：{value_text(scene.get('asset_assignment'))}",
            f"母素材/源长/时间线/播放模式：{value_text(scene.get('master_media_ids'))} / {value_text(scene.get('source_duration_seconds'))} / {value_text(scene.get('timeline_duration_seconds'))} / {value_text(scene.get('playback_mode'))}",
            f"原画保护/画幅/文字覆盖：{value_text(scene.get('media_body_blur_px'))} / {value_text(scene.get('frame_treatment'))} / {value_text(scene.get('text_overlay_engine'))}",
            f"双栏补位：{value_text(scene.get('argument_bridge'))}",
            f"动效：{value_text(scene.get('motion'))}",
            f"声音：{value_text(scene.get('sound'))}",
            f"转场：{value_text(scene.get('transition_in'))}",
            f"复用状态：{value_text(scene.get('reuse'))}",
            f"注意力重置：{value_text(scene.get('attention_reset'))}",
            "",
        ])
    lines.extend([
        "六、统一执行规则",
        "解释因果→动态图表/结构关系；提供证据→史料/权威截图；提供氛围→环境素材；强调转折→结构变化。抽象概念转为图形关系。",
        "连续多个年份、人物、制度名称或多层因果时加入时间轴、地图、结构图、数据变化或权力关系图。",
        "超过 30 分钟时每 5–8 分钟加入新问题、新案例、新视觉形式或新观点。",
        "同一母素材最多 3 次，强动作/人物/独特构图最多 2 次；每次复用至少改变 3 个维度且至少 2 个是视觉。",
        "同一母图片或母视频不得在相邻 Scene 连续出现；视频不得循环两遍填时长。跨后续章节复用仍执行间隔、三维变化、换前后镜和新增语义。",
        "严格按用户核对素材清单的顺序和语义段使用，不重新匹配、交换或跨段借用。图片展示5–10秒；视频>=5秒完整自然播放、<5秒播一次后保持末帧到5秒、>10秒不裁切，全部静音且不循环。",
        "图片/视频/证据图保持清晰原画全屏且blur=0；毛玻璃只用于上方文字。archive_evidence与book_evidence采用简中+英文来源区和大幅清晰证据图，在屏时禁止动态说理。",
        "底部固定保留220px字幕排除区；正文/卡片42–54px，辅助至少32px；右侧信息面竖向排列。",
        "同一总结/观点/程序化页面不得相邻连播或重复两遍补时长；visual_page_id或content_fingerprint相同直接失败。延长时使用一个HyperFrames段内至少3个新状态，否则换内容和版式都不同的新页。",
        "章节开头、因果、流程、制度关系、结构关系和文字说理必须程序化解释。锁定原文后先做video-shotcraft语义pass，并与RVE/Scenes/Curvable/Playground/HyperFrames候选比较；不设引擎占比或每章配额。",
        "argument_bridge只在缺少合适媒体或具体动态图解时使用：左双语大论点、右2–4个双语动态细分论点；目标4%–6%、硬上限6%、每章最多一场且不连续。",
        "章节接缝不得裸硬切；普通硬切只保留给有理由的语义撞击。右上角不放历史资料/情景演绎/AI生成等常驻标签。",
        "按锁定音频每300秒以内分段渲染，GPU if-possible；先运行 build_render_plan.py 探测硬件并生成自适应并发与显式缓存预算，再为当前段启动 watch_render_progress.py。每15秒采样资源和进度，无进展也每300秒固定心跳；180秒停滞自动诊断，300秒硬停滞立即报告，只处理当前段。CPU目标75–95%，GPU适用场景65–95%；RAM工作区间55–78%/硬上限85%，VRAM工作区间55–85%/硬上限92%。主录音优先，AI轻音乐或YouTube Studio Audio Library只交独立音频与审片混音。",
        "",
        "七、终检",
        "内容是否解释原因、是否逻辑跳跃、结论是否有依据；前 30 秒是否吸引、中段是否疲劳、结尾是否闭环；视觉是否服务内容、多语言是否易扩展。",
        "最终静音 MP4 必须通过 scripts/verify_master_media.py；QA 同时复核实际颜色、拉伸、黑边、圆窗、章节总览和画面语义，不能只看 metadata。审片混音另做三端试听。",
        "目标是理解效率最高，不是信息最多。",
        "",
        "八、结构化真本（不得省略）",
    ])
    payloads = {
        "project.json": project,
        "outline-scene-types.json": outline,
        "video-spec.json": spec,
        "animation-plan.json": animation,
        "asset-requirements.json": assets,
        "presenter-plan.json": presenter,
        "text-review.json": review,
    }
    for name, path in {
        "sound-plan.json": root / "08_audio" / "sound-plan.json",
        "render-plan.json": root / "09_qa" / "render-plan.json",
    }.items():
        if path.exists():
            payloads[name] = json.loads(path.read_text(encoding="utf-8-sig"))
        else:
            payloads[name] = {"status": "[待生成]", "expected_path": str(path)}
    for name, payload in payloads.items():
        lines.extend([f"--- {name} ---", json.dumps(payload, ensure_ascii=False, indent=2, sort_keys=True), ""])
    lines.append(f"项目根目录：{root}")
    return "\n".join(lines).rstrip() + "\n"


def build_image_research_doc(project: str, rows: list[dict[str, Any]], generated: datetime) -> Document:
    doc = Document()
    style_document(doc, "图片资料搜集需求单", project, generated)
    doc.add_paragraph("研究层先核验维基/Wikimedia Commons、公开数据库、博物馆、档案馆、图书馆与大学馆藏；资料采集 Skill 只按文字关键词和元数据检索，不识别图片内容。")
    if not rows:
        doc.add_paragraph("当前没有需要检索的真实图片。")
    for row in rows:
        add_entry(doc, f"{row.get('id', '?')}  |  {row.get('title', '未命名')}", [
            ("时间码 / 原句", f"{value_text(row.get('timecode'))} / {value_text(row.get('source_line'))}"),
            ("镜头功能", row.get("function")),
            ("可见对象", row.get("visible_subject")),
            ("中文检索词", row.get("query_zh")),
            ("英文检索词", row.get("query_en")),
            ("来源优先级", row.get("source_priority")),
            ("目标机构 / 数据库", row.get("target_sources")),
            ("年代 / 地点 / 人物", row.get("historical_scope")),
            ("动作 / 证据角色", f"{value_text(row.get('action'))} / {value_text(row.get('evidence_role'))}"),
            ("构图与安全区", row.get("composition")),
            ("原始比例 / 画幅处理", f"{value_text(row.get('source_aspect_ratio'))} / {value_text(row.get('frame_treatment'))}"),
            ("清晰原图全屏动态 / 局部文字毛玻璃", f"{value_text(row.get('image_motion'))} / {value_text(row.get('glass_problem_points'))}"),
            ("用户锁定顺序 / 语义段", f"{value_text(row.get('manifest_order'))} / {value_text(row.get('assigned_script_segment_id'))}"),
            ("历史资料标题来源区 / 大幅证据图", row.get("source_file_metadata")),
            ("最低像素", row.get("minimum_pixels")),
            ("史实核验点", row.get("verification")),
            ("授权 / 署名", row.get("license_requirements")),
            ("候选 URL", row.get("candidate_urls")),
            ("文字检索状态", row.get("search_status")),
            ("用户核对 / 指派锁", f"{value_text(row.get('user_verified'))} / {value_text(row.get('assignment_locked'))}"),
            ("备注", row.get("notes")),
        ])
    return doc


def build_ai_image_doc(project: str, rows: list[dict[str, Any]], generated: datetime) -> Document:
    doc = Document()
    style_document(doc, "图片AI生成需求单", project, generated)
    doc.add_paragraph("图片数量不设上限，按锁定文案时间线与5–10秒计划展示时长覆盖视觉缺口；长段落拆成主画面、细节、状态变化、衔接和备用。所有提示词须写实、专业、可执行；准确文字、数字、地图边界和史料原文由后期叠加。archive_evidence与book_evidence不以AI图冒充证据。")
    if not rows:
        doc.add_paragraph("当前没有需要 AI 生成的图片。")
    for row in rows:
        add_entry(doc, f"{row.get('id', '?')}  |  {row.get('title', '未命名')}", [
            ("语义段 / 原文 / 时间码", f"{value_text(row.get('semantic_segment_id'))} / {value_text(row.get('source_line'))} / {value_text(row.get('timecode'))}"),
            ("用途 / 计划展示秒数", f"{value_text(row.get('function'))} / {value_text(row.get('planned_duration_seconds'))}"),
            ("变体角色 / 预期使用次数", f"{value_text(row.get('variant_role'))} / {value_text(row.get('expected_use_count', 1))}"),
            ("主体 / 动作 / 环境", f"{value_text(row.get('subject'))} / {value_text(row.get('action'))} / {value_text(row.get('environment'))}"),
            ("时代 / 地点", row.get("setting")),
            ("镜头 / 焦段 / 构图安全区", row.get("camera_composition")),
            ("光线 / 材质 / 现实瑕疵", row.get("realism")),
            ("连续性锚点", row.get("continuity_anchor")),
            ("比例 / 分辨率", row.get("format")),
            ("画幅处理", row.get("frame_treatment")),
            ("中文提示词", row.get("prompt_zh")),
            ("English prompt", row.get("prompt_en")),
            ("文字后期叠加", row.get("post_text")),
            ("审核状态", row.get("status")),
            ("拟分配语义段 / 用户核对", f"{value_text(row.get('assigned_script_segment_id'))} / {value_text(row.get('user_verified'))}"),
        ])
    return doc


def build_ai_video_doc(project: str, rows: list[dict[str, Any]], generated: datetime) -> Document:
    doc = Document()
    style_document(doc, "视频AI生成需求单", project, generated)
    doc.add_paragraph(f"共 {len(rows)} 条（上限 250）；每条按语义选择5–15秒、16:9、无声、无字幕、写实、一个主体动作和一个克制镜头动作，并保留稳定结尾。进入成片后完整播放，原画全屏、零模糊、静音且不循环；视频在屏时减少额外动态说理。")
    if not rows:
        doc.add_paragraph("当前没有需要 AI 生成的视频。")
    for row in rows:
        add_entry(doc, f"{row.get('id', '?')}  |  {row.get('title', '未命名')}  |  {value_text(row.get('duration_seconds'))}秒", [
            ("成片时间码 / 功能", f"{value_text(row.get('timecode'))} / {value_text(row.get('function'))}"),
            ("唯一主体动作", row.get("subject_action")),
            ("唯一镜头动作", row.get("camera_motion")),
            ("时代 / 地点 / 环境", row.get("setting")),
            ("连续性", row.get("continuity_anchor")),
            ("写实光线与材质", row.get("realism")),
            ("安全区", row.get("safe_zone")),
            ("开头 / 结尾状态", f"{value_text(row.get('opening_state'))} / {value_text(row.get('ending_state'))}"),
            ("目标规格", row.get("format", "5–15s / 16:9 / 2560×1440 / 60fps / silent")),
            ("字幕要求", "不生成任何字幕、标题、水印或可读文字"),
            ("中文提示词", row.get("prompt_zh")),
            ("English prompt", row.get("prompt_en")),
            ("验收点", row.get("acceptance")),
            ("拟分配语义段 / 用户核对", f"{value_text(row.get('assigned_script_segment_id'))} / {value_text(row.get('user_verified'))}"),
        ])
    return doc


def build_presenter_doc(
    project: str,
    presenter: dict[str, Any],
    total_duration: float,
    generated: datetime,
) -> Document:
    rows = presenter.get("segments", [])
    presenter_seconds = 0.0
    for row in rows:
        try:
            presenter_seconds += parse_timecode(row.get("end")) - parse_timecode(row.get("start"))
        except ValueError:
            pass
    ratio = presenter_seconds / total_duration if total_duration > 0 else 0
    doc = Document()
    style_document(doc, "真人露脸补录台词清单", project, generated)
    doc.sections[0].top_margin = Cm(1.7)
    doc.sections[0].bottom_margin = Cm(1.3)
    doc.styles["Normal"].font.size = Pt(10)
    add_label(doc, "源视频原则", "统一横屏 16:9 全屏拍摄、人物居中、保留双手和左右内容感知圆形头肩裁切安全区；后期裁成直径约360px并抬到字幕排除区之上的圆窗。")
    add_label(doc, "录制流程", "前 5 秒纯静音中性姿态 → 双短提示音 → 连续完整台词 → 双短提示音 → 后 5 秒纯静音中性姿态。")
    add_label(doc, "成片预算", f"{presenter_seconds:.1f} 秒 / {total_duration:.1f} 秒 = {ratio:.2%}；必须 <20%。")
    add_label(doc, "版式策略", "只有总开头和实际最终结尾全屏；正文固定使用抬高右下圆窗（340–380px，基准360px）；最终全屏持续到最后一帧。说明动效逐镜比较Remotion与HyperFrames。")
    if not rows:
        doc.add_paragraph("当前没有真人补录段。")
    for row in rows:
        add_entry(doc, f"{row.get('id', '?')}  |  {row.get('chapter', '未分章')}  |  {row.get('final_mode', '抬高内容感知圆窗')}", [
            ("成片时间 / SRT", f"{value_text(row.get('start'))}–{value_text(row.get('end'))} / {value_text(row.get('srt_lines'))}"),
            ("内容位置", row.get("content_position")),
            ("源视频拍法", row.get("source_capture", "16:9 全屏、人物居中、圆形头肩裁切安全、60fps、固定曝光/白平衡/焦点")),
            ("文件名", row.get("filename")),
            ("完整起句", row.get("opening_sentence")),
            ("完整收句", row.get("closing_sentence")),
            ("连续完整台词", row.get("full_script")),
            ("手势点", row.get("gesture")),
            ("后期处理", row.get("post_process")),
            ("精确原声切片", row.get("exact_audio_file", f"{row.get('id', 'R-xx')}_exact.wav")),
            ("5秒空首尾提示音版", row.get("guide_audio_file", f"{row.get('id', 'R-xx')}_5s-beep-guide.wav")),
        ])
    doc.add_heading("交付检查", level=1)
    for item in [
        "每条从完整句开始到完整句结束；没有半句话或过渡词单独成段。",
        "源视频均为独立全屏文件，编号不跳号；服装、机位、焦距、光位、曝光和白平衡一致。",
        "人物头肩和口型始终位于中央圆形裁切安全区；后期可无损裁成左右择位且抬高的圆窗；关键手势不依赖圆窗展示。",
        "每条前后各 5 秒纯静音，起止均有双短提示音；对应两类音频文件存在。",
        "成片真人累计占比严格小于 20%。",
    ]:
        doc.add_paragraph(item, style="List Number")
    return doc


def main() -> int:
    parser = argparse.ArgumentParser(description="Build all jian-ji-liu-cheng planning documents.")
    parser.add_argument("project_root", type=Path)
    args = parser.parse_args()
    root = args.project_root.expanduser().resolve()
    project = load_json(root / "project.json")
    outline = load_json(root / "04_spec" / "outline-scene-types.json")
    spec = load_json(root / "04_spec" / "video-spec.json")
    animation = load_json(root / "04_spec" / "animation-plan.json")
    assets = load_json(root / "04_spec" / "asset-requirements.json")
    presenter = load_json(root / "05_assets" / "presenter" / "presenter-plan.json")
    review = load_json(root / "03_review" / "text-review.json")
    brief_path = root / "01_brief" / "project-brief.md"
    brief = brief_path.read_text(encoding="utf-8-sig") if brief_path.exists() else "[待补充]"
    errors = validate_inputs(project, spec, assets)
    if errors:
        raise SystemExit("Invalid delivery inputs:\n- " + "\n- ".join(errors))

    generated = datetime.now().astimezone()
    out = root / "10_deliverables"
    out.mkdir(parents=True, exist_ok=True)
    project_name = str(project.get("project", spec.get("project", "未命名项目")))
    manual = build_execution_manual(root, project, brief, outline, spec, animation, assets, presenter, review, generated)
    manual_path = out / "视频剪辑执行说明书.txt"
    manual_path.write_text(manual, encoding="utf-8-sig")
    documents = {
        "图片资料搜集需求单.docx": build_image_research_doc(
            project_name, assets.get("real_image_requests", []), generated
        ),
        "图片AI生成需求单.docx": build_ai_image_doc(
            project_name, assets.get("ai_image_requests", []), generated
        ),
        "视频AI生成需求单.docx": build_ai_video_doc(
            project_name, assets.get("ai_video_requests", []), generated
        ),
        "真人露脸补录台词清单.docx": build_presenter_doc(
            project_name,
            presenter,
            float(spec.get("audio_duration_seconds") or project.get("audio_duration_seconds") or 0),
            generated,
        ),
        "动效文字执行手册.docx": build_text_review_document(review, generated),
    }
    saved = [str(manual_path)]
    for name, doc in documents.items():
        path = out / name
        save_document(doc, path)
        saved.append(str(path))
    review_target = root / "03_review" / "动效文字执行手册.docx"
    shutil.copy2(out / "动效文字执行手册.docx", review_target)

    result = {
        "project_root": str(root),
        "generated_at": generated.isoformat(timespec="seconds"),
        "documents": saved,
        "text_review_unresolved": unresolved_rows(review),
        "ai_video_count": len(assets.get("ai_video_requests", [])),
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
