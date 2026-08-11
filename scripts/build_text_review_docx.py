#!/usr/bin/env python3
"""Build the timestamped motion-text review manual from JSON."""

from __future__ import annotations

import argparse
import json
import shutil
import sys
from collections import Counter
from datetime import datetime
from pathlib import Path

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except ImportError as exc:  # pragma: no cover
    raise SystemExit("python-docx is required. Use the bundled Codex document Python runtime.") from exc


STATUSES = {"保留", "排版调整", "建议修改", "必须修改", "待确认", "已确认"}
BLOCKING = {"必须修改", "待确认"}
CHANGE_KINDS = {"unchanged_approved", "layout_only", "new", "changed"}


def local_now() -> datetime:
    return datetime.now().astimezone()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_doc_fonts(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Source Sans 3"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(61, 48, 40)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Serif SC")
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[name]
        style.font.name = "Lora"
        style.font.color.rgb = RGBColor(61, 48, 40) if name != "Heading 2" else RGBColor(232, 115, 74)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Serif SC")


def add_kv_table(doc: Document, items: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for key, value in items:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value or "—"
        set_cell_shading(cells[0], "FFF5EA")
        for run in cells[0].paragraphs[0].runs:
            run.bold = True


def validate_payload(data: dict) -> list[str]:
    errors: list[str] = []
    if not data.get("project"):
        errors.append("project is required")
    if data.get("review_mode") != "delta_only":
        errors.append("review_mode must be delta_only")
    if data.get("baseline_existing_text_approved") is not True:
        errors.append("baseline_existing_text_approved must be true")
    if not isinstance(data.get("rows"), list):
        errors.append("rows must be a list")
        return errors
    seen: set[str] = set()
    for index, row in enumerate(data["rows"], 1):
        rid = str(row.get("id", "")).strip()
        if not rid:
            errors.append(f"row {index}: id is required")
        elif rid in seen:
            errors.append(f"row {index}: duplicate id {rid}")
        seen.add(rid)
        if not row.get("start") or not row.get("end"):
            errors.append(f"row {index}: start and end are required")
        status = row.get("status", "待确认")
        if status not in STATUSES:
            errors.append(f"row {index}: unsupported status {status}")
        if row.get("change_kind") not in CHANGE_KINDS:
            errors.append(f"row {index}: unsupported change_kind {row.get('change_kind')!r}")
    return errors


def unresolved_rows(data: dict) -> list[str]:
    unresolved: list[str] = []
    for row in data.get("rows", []):
        if row.get("change_kind") not in {"new", "changed"}:
            continue
        rid = str(row.get("id", "?"))
        if row.get("status", "待确认") in BLOCKING:
            unresolved.append(rid)
        elif not row.get("approval_version") or not row.get("approval_source"):
            unresolved.append(rid)
    return unresolved


def build_document(data: dict, generated: datetime) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    set_doc_fonts(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("动效文字执行手册")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(232, 115, 74)

    generated_text = generated.isoformat(timespec="seconds")
    add_kv_table(doc, [
        ("项目", str(data.get("project", ""))),
        ("版本", str(data.get("version", "V1"))),
        ("生成时间", generated_text),
        ("帧率", f"{data.get('fps', 60)} fps"),
        ("审阅模式", "仅新增/改字（delta_only）"),
        ("既有文字基线", "已全部通过" if data.get("baseline_existing_text_approved") else "未确认"),
        ("源文件", "、".join(map(str, data.get("source_files", []))) or "未列出"),
    ])

    doc.add_heading("审阅规则", level=1)
    doc.add_paragraph("既有文字基线已全部通过；只有 change_kind=new/changed 的增量行需要审阅。增量行中的“必须修改”“待确认”或缺少用户批准证据会阻塞对应含文字画面；layout_only 不重审，任何实际改字都必须更新版本和本表。")

    counts = Counter(str(row.get("status", "待确认")) for row in data.get("rows", []))
    blockers = unresolved_rows(data)
    doc.add_heading("状态汇总", level=1)
    summary = "；".join(f"{status} {counts.get(status, 0)}" for status in ["保留", "已确认", "排版调整", "建议修改", "必须修改", "待确认"])
    doc.add_paragraph(summary)
    doc.add_paragraph("阻塞项：" + ("、".join(blockers) if blockers else "无"))

    doc.add_heading("逐条审阅", level=1)
    if not data.get("rows"):
        doc.add_paragraph("当前没有新增或改字条目；既有文字继承已通过基线，不重复送审。")

    fields = [
        ("章节 / 载体", "{chapter} / {carrier}"),
        ("原文", "{source_zh}"),
        ("英文主文", "{screen_en}"),
        ("中文标注", "{screen_zh}"),
        ("内容类型", "{content_type}"),
        ("改动类型", "{change_kind}"),
        ("当前问题", "{issue}"),
        ("修改建议", "{suggestion}"),
        ("修改理由", "{reason}"),
        ("用户决定", "{user_decision}"),
        ("批准版本", "{approval_version}"),
        ("批准日期", "{approved_at}"),
        ("批准来源", "{approval_source}"),
    ]
    for row in data.get("rows", []):
        status = str(row.get("status", "待确认"))
        heading = f"{row.get('id', '')}  |  {row.get('start', '')} – {row.get('end', '')}  |  {status}"
        doc.add_heading(heading, level=2)
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for label, template in fields:
            cells = table.add_row().cells
            cells[0].text = label
            try:
                value = template.format(**{k: str(v or "") for k, v in row.items()})
            except KeyError:
                value = ""
            cells[1].text = value or "—"
            set_cell_shading(cells[0], "FFF5EA")
            for run in cells[0].paragraphs[0].runs:
                run.bold = True
        doc.add_paragraph()

    doc.add_heading("用户签批", level=1)
    add_kv_table(doc, [("审阅版本", ""), ("日期", ""), ("结论", "放行 / 修改后再审 / 暂停"), ("备注", "")])
    return doc


def main() -> int:
    parser = argparse.ArgumentParser(description="Build the timestamped motion-text review Word gate.")
    parser.add_argument("input_json", type=Path)
    parser.add_argument("output_docx", type=Path, help="Stable latest output path")
    parser.add_argument("--no-archive", action="store_true", help="Do not create a versioned archive copy")
    args = parser.parse_args()

    data = json.loads(args.input_json.read_text(encoding="utf-8-sig"))
    errors = validate_payload(data)
    if errors:
        raise SystemExit("Invalid review JSON:\n- " + "\n- ".join(errors))

    generated = local_now()
    args.output_docx.parent.mkdir(parents=True, exist_ok=True)
    version = str(data.get("version", "V1")).replace(" ", "_")
    stamp = generated.strftime("%Y%m%d-%H%M%S")
    archive = args.output_docx.with_name(f"{args.output_docx.stem}_{version}_{stamp}{args.output_docx.suffix}")
    doc = build_document(data, generated)
    target = args.output_docx if args.no_archive else archive
    doc.save(target)
    if not args.no_archive:
        shutil.copy2(archive, args.output_docx)

    blockers = unresolved_rows(data)
    print(json.dumps({
        "latest": str(args.output_docx.resolve()),
        "archive": None if args.no_archive else str(archive.resolve()),
        "generated_at": generated.isoformat(timespec="seconds"),
        "rows": len(data.get("rows", [])),
        "blocking_items": blockers,
        "released": data.get("baseline_existing_text_approved") is True and not blockers,
    }, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
